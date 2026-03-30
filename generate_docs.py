from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"c:\Users\jflin\WorkBuddy\20260329125422\stocktrading"

# ─── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_table(table, header_color="1F3864"):
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after  = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold  = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                set_cell_bg(cell, header_color)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 0:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(18)
    elif level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(14)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(12)
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(11)
    return p

def add_meta(doc, version, date, extra=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"版本：{version}    日期：{date}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    if extra:
        r2 = p.add_run(f"    {extra}")
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph()

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

# ─── Doc 1: User Requirements ───────────────────────────────────────────────

def build_doc1():
    doc = Document()
    set_doc_margins(doc)

    add_heading(doc, "香港股票投资分析 App — 用户需求文档", 0)
    add_meta(doc, "v1.0", "2026-03-29", "平台：Android（可下载安装）")

    # 1. 项目概述
    add_heading(doc, "1. 项目概述", 1)
    add_body(doc, "本应用旨在帮助个人投资者分析香港股票市场（港股 / HKEX），追踪行情走势，记录个人买卖交易，并提供每日行情摘要，辅助做出更理性的投资决策。")

    # 2. 目标用户
    add_heading(doc, "2. 目标用户", 1)
    t = doc.add_table(rows=5, cols=2)
    rows_data = [
        ("用户类型", "描述"),
        ("主要用户", "港股个人散户投资者"),
        ("使用频率", "每个交易日（每日使用）"),
        ("技术程度", "普通手机用户，不需要专业编程知识"),
        ("语言偏好", "繁体中文为主，英文辅助"),
    ]
    for i, (a, b) in enumerate(rows_data):
        t.cell(i, 0).text = a
        t.cell(i, 1).text = b
    style_table(t)

    # 3. 功能需求
    add_heading(doc, "3. 功能需求（Functional Requirements）", 1)

    sections = [
        ("3.1 市场概览（Market Overview）", [
            "显示恒生指数（HSI）、国企指数（HSCEI）、恒生科技指数（HSTECH）实时数据",
            "显示港股市场当日涨跌幅排行（升幅榜 / 跌幅榜）",
            "显示成交额最高股票排行",
            "显示大盘走势图（日 / 周 / 月 / 年）",
        ]),
        ("3.2 行业分析（Sector Analysis）", [
            "港股主要行业分类（地产、金融、科技、医疗、能源、消费品等）",
            "各行业当日涨跌表现（行业热力图）",
            "行业内龙头股列表及走势",
            "行业资金流向（净买入 / 净卖出）",
        ]),
        ("3.4 自选股票组合（Watchlist）", [
            "用户可添加 / 删除自选股票",
            "自选股票实时报价一览",
            "支持分组管理（如：科技股、蓝筹股、待观察等）",
        ]),
        ("3.6 每日行情摘要（Daily Digest）", [
            "每个交易日收市后推送通知",
            "大盘表现总结",
            "自选股票涨跌情况",
            "持仓股当日表现",
            "重要财经新闻摘要",
            "用户可设定推送时间（默认：港股收市后 16:30 HKT）",
        ]),
        ("3.7 投资分析工具（Analysis Tools）", [
            "股票对比工具（最多 3 只股票同时对比走势）",
            "仓位计算器（根据止损位计算建议仓位大小）",
            "股息计算器（预测年化股息收入）",
        ]),
        ("3.8 设置与账户（Settings）", [
            "语言设置（繁体中文 / 简体中文 / English）",
            "货币显示（港元 HKD）",
            "推送通知开关及时间设定",
            "数据刷新频率设定（15秒 / 1分钟 / 5分钟）",
            "本地数据备份 / 导出（CSV 格式）",
            "深色模式 / 浅色模式",
        ]),
    ]
    for title, bullets in sections:
        add_heading(doc, title, 2)
        for b in bullets:
            add_bullet(doc, b)

    # 3.3 个股分析 — special
    add_heading(doc, "3.3 个股分析（Individual Stock Analysis）", 2)
    add_bullet(doc, "股票搜索（名称 / 股票代码）")
    add_heading(doc, "股票基本资料", 3)
    for b in ["公司名称、股票代码、所属行业", "实时价格、涨跌幅、成交量、市值"]:
        add_bullet(doc, b, 1)
    add_heading(doc, "技术分析图表", 3)
    for b in ["K 线图（日 K / 周 K / 月 K）",
              "常用技术指标：MA（5/10/20/50/200日）、MACD、RSI、布林带（Bollinger Bands）",
              "支撑位 / 阻力位标注"]:
        add_bullet(doc, b, 1)
    add_heading(doc, "基本面数据", 3)
    for b in ["市盈率（P/E）、市净率（P/B）、股息率", "52 周高低位", "近期财报摘要"]:
        add_bullet(doc, b, 1)
    add_bullet(doc, "AI 走势分析摘要（简单买入 / 持有 / 卖出建议，附理由）")
    add_bullet(doc, "相关新闻列表（来源：财经新闻 API）")

    # 3.5 交易记录
    add_heading(doc, "3.5 交易记录（Trade Journal）", 2)
    add_heading(doc, "手动记录每笔买卖", 3)
    for b in ["股票代码 / 名称", "买入 / 卖出方向", "价格、数量、手续费", "日期 / 时间", "备注（如入场理由）"]:
        add_bullet(doc, b, 1)
    add_heading(doc, "持仓总览", 3)
    for b in ["当前持仓列表、平均买入价、现价、浮动盈亏（金额 + 百分比）",
              "总投入资金、总市值、总盈亏"]:
        add_bullet(doc, b, 1)
    add_bullet(doc, "已平仓记录（历史交易清单）")
    add_bullet(doc, "盈亏统计图表（月度 / 年度）")

    # 4. 非功能需求
    add_heading(doc, "4. 非功能需求（Non-Functional Requirements）", 1)
    nfr = [
        ("类别", "要求"),
        ("性能", "实时行情数据延迟 ≤ 15 秒（免费方案）"),
        ("可用性", "界面简洁直观，3 步内完成常用操作"),
        ("稳定性", "崩溃率 < 1%，网络断线时可查看离线缓存"),
        ("安全性", "交易记录加密存储，无需登录账户（本地存储为主）"),
        ("兼容性", "Android 8.0 及以上版本"),
        ("本地化", "繁体中文界面，港股时区（HKT UTC+8）"),
    ]
    t2 = doc.add_table(rows=len(nfr), cols=2)
    for i, (a, b) in enumerate(nfr):
        t2.cell(i, 0).text = a
        t2.cell(i, 1).text = b
    style_table(t2)

    # 5. 数据来源
    add_heading(doc, "5. 数据来源需求", 1)
    datasrc = [
        ("数据类型", "建议来源"),
        ("实时 / 延迟报价", "Yahoo Finance API、Alpha Vantage、Futu OpenAPI"),
        ("历史价格数据", "Yahoo Finance（yfinance Python库）"),
        ("财经新闻", "NewsAPI、AASTOCKS、信报财经 RSS"),
        ("基本面数据", "Yahoo Finance、Macrotrends"),
        ("行业分类", "HKEX 官网数据"),
    ]
    t3 = doc.add_table(rows=len(datasrc), cols=2)
    for i, (a, b) in enumerate(datasrc):
        t3.cell(i, 0).text = a
        t3.cell(i, 1).text = b
    style_table(t3)

    # 6. 界面设计
    add_heading(doc, "6. 界面设计要求", 1)
    ui_items = [
        "风格：现代金融 App 风格（参考老虎证券、富途牛牛）",
        "主色调：深色主题为主（深蓝 / 黑色背景），红色表示上涨，绿色表示下跌（港股惯例）",
        "图表库：支持流畅滑动的 K 线图及指标图",
        "底部导航栏：市场 / 自选 / 交易记录 / 分析工具 / 设置",
    ]
    for item in ui_items:
        add_bullet(doc, item)

    # 7. V2.0
    add_heading(doc, "7. 未来版本功能（V2.0 规划）", 1)
    v2 = [
        "连接港股券商 API 实现自动交易",
        "AI 选股推荐（基于技术面 + 基本面综合评分）",
        "组合风险评估（Beta 值、夏普比率）",
        "社群讨论功能",
        "iPad / iOS 版本",
    ]
    for item in v2:
        add_bullet(doc, item)

    path = os.path.join(BASE, "01_User_Requirements.docx")
    doc.save(path)
    print(f"Saved: {path}")

# ─── Doc 2: Resources & Tech Stack ──────────────────────────────────────────

def build_doc2():
    doc = Document()
    set_doc_margins(doc)

    add_heading(doc, "香港股票投资分析 App — 所需资源及技术栈", 0)
    add_meta(doc, "v1.0", "2026-03-29")

    # 1. 架构
    add_heading(doc, "1. 技术架构总览", 1)
    arch_rows = [
        ("层次", "技术方案"),
        ("前端（Android App）", "Flutter 3.x（Dart 语言）"),
        ("通信协议", "REST API / WebSocket"),
        ("后端服务", "Python FastAPI"),
        ("数据层", "外部 API + 本地 SQLite"),
    ]
    t = doc.add_table(rows=len(arch_rows), cols=2)
    for i, (a, b) in enumerate(arch_rows):
        t.cell(i, 0).text = a
        t.cell(i, 1).text = b
    style_table(t)
    doc.add_paragraph()
    add_body(doc, "架构说明：前端优先 Flutter（Google 官方跨平台框架，性能好，适合金融图表）；后端使用 Python FastAPI（数据处理能力强，金融库生态丰富）；本地数据用 SQLite（交易记录、自选股等离线数据）。")

    # 2. 前端
    add_heading(doc, "2. 前端技术栈（Android App）", 1)
    fe = [
        ("技术", "用途", "推荐方案"),
        ("开发框架", "跨平台 Android App", "Flutter 3.x（Dart 语言）"),
        ("K 线图表", "股票走势图、技术指标", "k_chart / flutter_candlesticks"),
        ("普通图表", "饼图、柱状图、折线图", "fl_chart"),
        ("状态管理", "全局数据状态", "Riverpod 或 Bloc"),
        ("本地数据库", "交易记录离线存储", "sqflite（SQLite for Flutter）"),
        ("网络请求", "API 调用", "dio"),
        ("推送通知", "每日行情摘要推送", "flutter_local_notifications"),
        ("本地缓存", "股票数据缓存", "Hive / shared_preferences"),
        ("深色模式", "主题切换", "Flutter ThemeData"),
    ]
    t2 = doc.add_table(rows=len(fe), cols=3)
    for i, row in enumerate(fe):
        for j, val in enumerate(row):
            t2.cell(i, j).text = val
    style_table(t2)

    # 3. 后端
    add_heading(doc, "3. 后端技术栈", 1)
    be = [
        ("技术", "用途", "说明"),
        ("Python 3.11+", "主要开发语言", "金融数据处理库丰富"),
        ("FastAPI", "REST API 框架", "高性能，自动生成 API 文档"),
        ("yfinance", "获取 Yahoo Finance 数据", "免费，支持港股（代码加 .HK）"),
        ("pandas", "数据处理", "K 线数据计算、统计分析"),
        ("ta-lib / pandas-ta", "技术指标计算", "MA、MACD、RSI、布林带等"),
        ("APScheduler", "定时任务", "每日数据更新、推送触发"),
        ("SQLite / PostgreSQL", "数据持久化", "开发用 SQLite，生产可升级 PostgreSQL"),
        ("Redis（可选）", "实时数据缓存", "减少 API 调用频率"),
        ("WebSocket", "实时价格推送", "使用 websockets 库"),
    ]
    t3 = doc.add_table(rows=len(be), cols=3)
    for i, row in enumerate(be):
        for j, val in enumerate(row):
            t3.cell(i, j).text = val
    style_table(t3)

    # 4. 数据 API
    add_heading(doc, "4. 数据 API 资源", 1)
    add_heading(doc, "4.1 免费方案（推荐入门）", 2)
    free_api = [
        ("API", "数据类型", "限制", "港股支持"),
        ("Yahoo Finance (yfinance)", "延迟报价、历史数据、基本面", "无官方限制（非正式 API）", "✅ 优秀（代码：0700.HK）"),
        ("Alpha Vantage", "延迟报价、技术指标", "25次/天（免费版）", "✅ 支持"),
        ("NewsAPI", "财经新闻", "100次/天（免费版）", "✅ 英文新闻"),
        ("HKEX 官网 RSS", "官方公告", "无限制", "✅ 官方来源"),
    ]
    t4 = doc.add_table(rows=len(free_api), cols=4)
    for i, row in enumerate(free_api):
        for j, val in enumerate(row):
            t4.cell(i, j).text = val
    style_table(t4)

    add_heading(doc, "4.2 付费方案（进阶）", 2)
    paid_api = [
        ("API", "月费", "优势"),
        ("Polygon.io", "~$29/月", "实时数据，稳定可靠"),
        ("Futu OpenAPI", "免费（需开户）", "专业港股数据，富途证券官方"),
        ("Tiger Brokers API", "免费（需开户）", "老虎证券官方，实时数据"),
        ("Bloomberg API", "企业级定价", "最专业，超出个人需求"),
    ]
    t5 = doc.add_table(rows=len(paid_api), cols=3)
    for i, row in enumerate(paid_api):
        for j, val in enumerate(row):
            t5.cell(i, j).text = val
    style_table(t5)
    add_body(doc, "建议：初期使用 yfinance 免费方案，后期可接入 Futu OpenAPI（富途）以获得实时数据。")

    # 5. 云端部署
    add_heading(doc, "5. 云端部署资源（后端托管）", 1)
    cloud = [
        ("方案", "月费", "适合阶段"),
        ("本地运行（开发测试）", "免费", "开发阶段"),
        ("Railway.app", "~$5/月", "小型部署，简单易用"),
        ("Render.com", "免费起步", "个人项目推荐"),
        ("腾讯云轻量应用服务器", "~¥50/月", "国内访问稳定"),
        ("AWS / Google Cloud", "按用量计费", "生产环境扩展"),
    ]
    t6 = doc.add_table(rows=len(cloud), cols=3)
    for i, row in enumerate(cloud):
        for j, val in enumerate(row):
            t6.cell(i, j).text = val
    style_table(t6)

    # 6. 开发工具
    add_heading(doc, "6. 开发工具", 1)
    tools = [
        ("工具", "用途"),
        ("Android Studio", "Flutter Android 开发、调试、打包"),
        ("VS Code", "后端 Python 开发"),
        ("Flutter SDK", "App 开发框架"),
        ("Postman", "API 测试"),
        ("Git + GitHub", "版本控制"),
        ("Figma（可选）", "UI 设计稿"),
        ("DBeaver", "SQLite / 数据库管理"),
    ]
    t7 = doc.add_table(rows=len(tools), cols=2)
    for i, row in enumerate(tools):
        for j, val in enumerate(row):
            t7.cell(i, j).text = val
    style_table(t7)

    # 7. 人力资源
    add_heading(doc, "7. 人力资源需求", 1)
    add_heading(doc, "个人独立开发（推荐方案）", 2)
    hr = [
        ("角色", "技能要求", "时间投入"),
        ("全栈开发者（你）", "Flutter + Python 基础", "3-6 个月兼职"),
    ]
    t8 = doc.add_table(rows=len(hr), cols=3)
    for i, row in enumerate(hr):
        for j, val in enumerate(row):
            t8.cell(i, j).text = val
    style_table(t8)

    # 8. 预算
    add_heading(doc, "8. 预算估算", 1)
    add_heading(doc, "最低成本方案（个人开发）", 2)
    budget1 = [
        ("项目", "费用"),
        ("域名", "~$10/年"),
        ("服务器（Render 免费版）", "$0"),
        ("Yahoo Finance API", "$0"),
        ("NewsAPI（免费版）", "$0"),
        ("Android 开发者账号（Google Play）", "$25 一次性"),
        ("合计", "约 $35 起步"),
    ]
    t9 = doc.add_table(rows=len(budget1), cols=2)
    for i, row in enumerate(budget1):
        for j, val in enumerate(row):
            t9.cell(i, j).text = val
    style_table(t9)

    add_heading(doc, "进阶方案", 2)
    budget2 = [
        ("项目", "月费"),
        ("腾讯云服务器", "~¥50/月"),
        ("Futu OpenAPI", "免费（需开富途账户）"),
        ("NewsAPI Pro", "~$50/月"),
        ("合计", "约 ¥50-300/月"),
    ]
    t10 = doc.add_table(rows=len(budget2), cols=2)
    for i, row in enumerate(budget2):
        for j, val in enumerate(row):
            t10.cell(i, j).text = val
    style_table(t10)

    # 9. 第三方账号
    add_heading(doc, "9. 第三方服务账号清单", 1)
    accounts = [
        "Google Play Developer Account（发布 App）",
        "Yahoo Finance（直接使用，无需注册）",
        "Alpha Vantage 免费 API Key",
        "NewsAPI 免费 API Key",
        "GitHub 账号（代码版本管理）",
        "Render.com / Railway.app 账号（服务器托管）",
        "富途证券账号（可选，获取 Futu OpenAPI 权限）",
    ]
    for acc in accounts:
        add_bullet(doc, "☐  " + acc)

    path = os.path.join(BASE, "02_Resources_TechStack.docx")
    doc.save(path)
    print(f"Saved: {path}")

# ─── Doc 3: Project Plan ─────────────────────────────────────────────────────

def build_doc3():
    doc = Document()
    set_doc_margins(doc)

    add_heading(doc, "香港股票投资分析 App — 项目计划", 0)
    add_meta(doc, "v1.0", "2026-03-29", "预计总工期：约 20 周（5 个月）")

    # 1. 目标
    add_heading(doc, "1. 项目目标", 1)
    add_body(doc, "在 5 个月内完成一个可在 Android 手机安装的港股投资辅助 App，具备市场分析、个股查询、交易记录管理及每日推送功能。")

    # 2. 阶段概览
    add_heading(doc, "2. 项目阶段概览", 1)
    phases = [
        ("阶段", "内容", "周期"),
        ("阶段 1", "规划与设计", "第 1-2 周"),
        ("阶段 2", "环境搭建", "第 3 周"),
        ("阶段 3", "后端开发", "第 4-8 周"),
        ("阶段 4", "前端开发", "第 9-16 周"),
        ("阶段 5", "测试与优化", "第 17-19 周"),
        ("阶段 6", "发布上线", "第 20 周"),
    ]
    tp = doc.add_table(rows=len(phases), cols=3)
    for i, row in enumerate(phases):
        for j, val in enumerate(row):
            tp.cell(i, j).text = val
    style_table(tp)

    # 3. 详细计划
    add_heading(doc, "3. 详细阶段计划", 1)

    # 阶段1
    add_heading(doc, "阶段 1：规划与设计（第 1-2 周）", 2)
    add_body(doc, "目标：确定功能范围，完成 UI 设计稿")
    ph1 = [
        ("任务", "详情", "产出物"),
        ("1.1 确认功能需求", "基于需求文档确认 MVP 功能范围", "最终需求文档"),
        ("1.2 数据库设计", "设计交易记录、自选股、缓存数据表结构", "数据库 ER 图"),
        ("1.3 API 设计", "定义前后端接口规范", "API 文档（OpenAPI）"),
        ("1.4 UI 线框图", "设计各主要页面布局", "Figma 线框图"),
        ("1.5 注册账号", "注册所有第三方服务账号", "账号清单"),
    ]
    t1 = doc.add_table(rows=len(ph1), cols=3)
    for i, row in enumerate(ph1):
        for j, val in enumerate(row):
            t1.cell(i, j).text = val
    style_table(t1)
    doc.add_paragraph()
    add_body(doc, "MVP 功能范围（第一版必须有）：大盘指数显示 · 个股搜索 + K 线图 · 自选股列表 · 交易记录（买卖记录 + 持仓）· 每日行情推送通知")

    # 阶段2
    add_heading(doc, "阶段 2：开发环境搭建（第 3 周）", 2)
    ph2 = [
        ("任务", "技术细节"),
        ("2.1 安装 Flutter SDK", "配置 Android 开发环境、模拟器"),
        ("2.2 安装 Python 环境", "Python 3.11 + FastAPI + 相关库"),
        ("2.3 初始化 Git 仓库", "创建 GitHub 仓库，建立分支策略"),
        ("2.4 搭建后端骨架", "FastAPI 项目结构，SQLite 初始化"),
        ("2.5 搭建前端骨架", "Flutter 项目结构，基础导航"),
        ("2.6 验证数据 API", "测试 yfinance 获取港股数据"),
    ]
    t2 = doc.add_table(rows=len(ph2), cols=2)
    for i, row in enumerate(ph2):
        for j, val in enumerate(row):
            t2.cell(i, j).text = val
    style_table(t2)

    # 阶段3
    add_heading(doc, "阶段 3：后端开发（第 4-8 周）", 2)
    backend_weeks = [
        ("第 4 周：数据服务层", [
            "封装 yfinance 数据获取模块",
            "实现港股报价查询接口（支持 .HK 代码）",
            "实现历史 K 线数据接口",
            "实现大盘指数（HSI / HSCEI / HSTECH）接口",
        ]),
        ("第 5 周：技术指标计算", [
            "集成 pandas-ta 库",
            "实现 MA（5/10/20/50/200）计算",
            "实现 MACD 计算",
            "实现 RSI 计算",
            "实现布林带（Bollinger Bands）计算",
        ]),
        ("第 6 周：行业 & 市场数据", [
            "实现行业分类接口",
            "实现涨跌幅排行榜接口",
            "实现成交量排行接口",
            "集成 NewsAPI 财经新闻接口",
        ]),
        ("第 7 周：交易记录 API", [
            "设计 SQLite 数据库表（trades, positions, watchlist）",
            "实现交易记录 CRUD 接口",
            "实现持仓计算逻辑（加权平均成本、浮动盈亏）",
            "实现盈亏统计接口",
        ]),
        ("第 8 周：推送 & 每日摘要", [
            "实现每日收市后数据汇总任务（APScheduler）",
            "实现推送触发接口",
            "后端 API 整体测试",
            "部署至 Render.com",
        ]),
    ]
    for week_title, tasks in backend_weeks:
        add_heading(doc, week_title, 3)
        for task in tasks:
            add_bullet(doc, "☐  " + task)

    # 阶段4
    add_heading(doc, "阶段 4：前端开发（第 9-16 周）", 2)
    frontend_weeks = [
        ("第 9-10 周：基础框架 & 市场概览页", [
            "搭建 Flutter 底部导航栏（5 个 Tab）",
            "实现深色主题 UI 系统",
            "实现市场概览页（大盘指数卡片）",
            "实现涨跌榜列表",
            "实现大盘走势折线图",
        ]),
        ("第 11-12 周：个股详情页", [
            "实现股票搜索功能",
            "实现 K 线图（日/周/月切换）",
            "实现技术指标叠加显示",
            "实现股票基本信息卡片",
            "实现相关新闻列表",
        ]),
        ("第 13 周：自选股页", [
            "实现自选股列表（实时报价）",
            "实现添加 / 删除自选股",
            "实现分组管理功能",
        ]),
        ("第 14-15 周：交易记录页", [
            "实现交易记录输入表单（买入 / 卖出）",
            "实现持仓总览（盈亏颜色显示）",
            "实现历史交易记录清单",
            "实现盈亏统计图表（月度柱状图）",
            "实现 CSV 导出功能",
        ]),
        ("第 16 周：分析工具 & 设置页", [
            "实现股票对比工具",
            "实现仓位计算器",
            "实现设置页（语言、推送、主题等）",
            "实现每日推送通知接收",
        ]),
    ]
    for week_title, tasks in frontend_weeks:
        add_heading(doc, week_title, 3)
        for task in tasks:
            add_bullet(doc, "☐  " + task)

    # 阶段5
    add_heading(doc, "阶段 5：测试与优化（第 17-19 周）", 2)
    tests = [
        ("测试类型", "内容"),
        ("功能测试", "逐一验证所有功能点"),
        ("性能测试", "图表加载速度、数据刷新流畅度"),
        ("兼容性测试", "Android 8.0 / 10 / 12 / 14 不同版本测试"),
        ("用户体验测试", "邀请 3-5 位真实用户试用，收集反馈"),
        ("网络异常测试", "弱网/断网时的降级处理"),
        ("数据准确性验证", "对比官方数据验证价格、指标计算正确性"),
    ]
    t5 = doc.add_table(rows=len(tests), cols=2)
    for i, row in enumerate(tests):
        for j, val in enumerate(row):
            t5.cell(i, j).text = val
    style_table(t5)
    doc.add_paragraph()
    add_heading(doc, "性能优化措施", 3)
    for item in ["K 线图数据分页加载（避免一次加载 5 年数据）", "实时数据 15 秒刷新，减少 API 调用", "图片/图表懒加载"]:
        add_bullet(doc, item)

    # 阶段6
    add_heading(doc, "阶段 6：发布上线（第 20 周）", 2)
    launch = [
        ("任务", "详情"),
        ("6.1 App 打包", "生成 Android APK / AAB 文件"),
        ("6.2 应用签名", "配置 Keystore 签名"),
        ("6.3 Google Play 上架", "准备截图、说明文字、隐私政策"),
        ("6.4 或直接 APK 分发", "通过邮件/链接直接安装（更快）"),
        ("6.5 监控上线", "观察崩溃报告、API 使用情况"),
    ]
    t6 = doc.add_table(rows=len(launch), cols=2)
    for i, row in enumerate(launch):
        for j, val in enumerate(row):
            t6.cell(i, j).text = val
    style_table(t6)

    # 4. 里程碑
    add_heading(doc, "4. 里程碑（Milestones）", 1)
    ms = [
        ("里程碑", "时间节点", "验收标准"),
        ("M1 - 需求确认", "第 2 周末", "需求文档签字确认"),
        ("M2 - 后端 MVP", "第 8 周末", "后端 API 全部可调用"),
        ("M3 - 前端 MVP", "第 16 周末", "完整 App 在模拟器上运行"),
        ("M4 - 测试完成", "第 19 周末", "测试通过率 ≥ 95%"),
        ("M5 - 正式上线", "第 20 周末", "APK 可安装使用"),
    ]
    t7 = doc.add_table(rows=len(ms), cols=3)
    for i, row in enumerate(ms):
        for j, val in enumerate(row):
            t7.cell(i, j).text = val
    style_table(t7)

    # 5. 风险
    add_heading(doc, "5. 风险管理", 1)
    risks = [
        ("风险", "可能性", "影响", "应对措施"),
        ("yfinance API 不稳定", "中", "高", "备选 Alpha Vantage / Futu API"),
        ("港股数据延迟 > 15 分钟", "低", "中", "使用付费实时 API（Futu）"),
        ("Flutter 图表库性能不足", "低", "高", "提前做 POC 验证"),
        ("开发时间超期", "中", "中", "优先保证 MVP 功能，次要功能延后"),
        ("Google Play 审核被拒", "低", "低", "先通过 APK 直接分发"),
    ]
    t8 = doc.add_table(rows=len(risks), cols=4)
    for i, row in enumerate(risks):
        for j, val in enumerate(row):
            t8.cell(i, j).text = val
    style_table(t8)

    # 6. 成功标准
    add_heading(doc, "6. 项目成功标准", 1)
    success = [
        "在 Android 手机上安装并启动 App",
        "查看港股大盘及个股实时（延迟）行情",
        "查看 K 线图及技术指标",
        "管理自选股列表",
        "记录每笔买卖，查看持仓盈亏",
        "每日收市后收到行情推送摘要",
    ]
    for item in success:
        add_bullet(doc, "☐  " + item)

    path = os.path.join(BASE, "03_Project_Plan.docx")
    doc.save(path)
    print(f"Saved: {path}")

# ─── main ────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_doc1()
    build_doc2()
    build_doc3()
    print("All 3 Word documents generated successfully!")
