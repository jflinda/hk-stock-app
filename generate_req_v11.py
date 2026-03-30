from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"c:\Users\jflin\WorkBuddy\20260329125422\stocktrading"

# ─── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def style_table(table, header_color="1F3864"):
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(3)
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold  = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                set_cell_bg(cell, header_color)

def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    if level == 0:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(18)
        run.font.bold = True
    elif level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.size = Pt(13)
        run.font.bold = True
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(11.5)
        run.font.bold = True
    else:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.italic = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_meta_block(doc, version, date, status):
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    cells = [
        f"Version: {version}",
        f"Date: {date}",
        f"Status: {status}",
    ]
    for i, txt in enumerate(cells):
        c = t.cell(0, i)
        c.text = txt
        set_cell_bg(c, "2E74B5")
        for para in c.paragraphs:
            for run in para.runs:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size  = Pt(10)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
    doc.add_paragraph()

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p

def add_note(doc, text):
    """Highlighted note box"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("📌  " + text)
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def set_doc_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ─── Build Document ──────────────────────────────────────────────────────────

def build_user_requirements():
    doc = Document()
    set_doc_margins(doc)

    # Title
    add_heading(doc, "Hong Kong Stock Investment Analysis App", 0)
    p = doc.add_paragraph()
    r = p.add_run("User Requirements Document")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    r.font.bold = True
    p.paragraph_format.space_after = Pt(6)

    add_meta_block(doc, "v1.1", "2026-03-29", "Approved for Design & Development")
    add_divider(doc)

    # ── 1. Project Overview ──
    add_heading(doc, "1. Project Overview", 1)
    add_body(doc, "This application is designed to help individual investors analyse the Hong Kong stock market (HKEX), track market trends, record personal buy/sell transactions, review portfolio performance, and receive daily market digests — empowering users to make more informed investment decisions.")
    add_divider(doc)

    # ── 2. Target Users ──
    add_heading(doc, "2. Target Users", 1)
    users = [
        ("Attribute", "Description"),
        ("Primary User", "Individual retail investors trading Hong Kong stocks"),
        ("Usage Frequency", "Every trading day (daily use)"),
        ("Technical Level", "General smartphone user; no programming knowledge required"),
        ("Default Language", "English (user may switch to Traditional Chinese)"),
    ]
    t = doc.add_table(rows=len(users), cols=2)
    for i, (a, b) in enumerate(users):
        t.cell(i, 0).text = a
        t.cell(i, 1).text = b
    style_table(t)
    add_divider(doc)

    # ── 3. Functional Requirements ──
    add_heading(doc, "3. Functional Requirements", 1)

    # 3.1
    add_heading(doc, "3.1  Market Overview", 2)
    for b in [
        "Display real-time (or delayed) data for major indices: Hang Seng Index (HSI), H-Shares Index (HSCEI), Hang Seng Tech Index (HSTECH)",
        "Display top gainers and top losers for the day (by % change)",
        "Display most active stocks by turnover value",
        "Display intraday and historical market trend charts (Daily / Weekly / Monthly / Yearly)",
    ]:
        add_bullet(doc, b)

    # 3.2
    add_heading(doc, "3.2  Sector Analysis", 2)
    for b in [
        "Browse Hong Kong stocks by major industry sectors (Property, Finance, Technology, Healthcare, Energy, Consumer Goods, etc.)",
        "View daily sector performance heatmap (sector-level gain/loss at a glance)",
        "View list of leading stocks within each sector and their trends",
        "View sector capital flow (net buy / net sell)",
    ]:
        add_bullet(doc, b)

    # 3.3
    add_heading(doc, "3.3  Individual Stock Analysis", 2)
    add_bullet(doc, "Search stocks by name or ticker code (e.g. 0700.HK)")
    add_heading(doc, "Stock Summary Card", 3)
    for b in ["Company name, ticker code, sector", "Real-time price, % change, volume, market capitalisation"]:
        add_bullet(doc, b, 1)
    add_heading(doc, "Technical Analysis Charts", 3)
    for b in [
        "Candlestick (K-Line) chart — Daily / Weekly / Monthly",
        "Technical indicators: MA (5/10/20/50/200-day), MACD, RSI, Bollinger Bands",
        "Support and resistance level annotations",
    ]:
        add_bullet(doc, b, 1)
    add_heading(doc, "Fundamental Data", 3)
    for b in ["P/E Ratio, P/B Ratio, Dividend Yield", "52-week high / low", "Recent earnings summary"]:
        add_bullet(doc, b, 1)
    add_bullet(doc, "AI-generated trend analysis summary (Buy / Hold / Sell recommendation with reasoning)")
    add_bullet(doc, "Related financial news list (sourced via financial news API)")

    # 3.4
    add_heading(doc, "3.4  Watchlist", 2)
    for b in [
        "Add / remove stocks to a personal watchlist",
        "View real-time quotes for all watchlist stocks on a single screen",
        "Support grouping / tagging (e.g. Blue Chips, Tech, Watch List, etc.)",
    ]:
        add_bullet(doc, b)

    # 3.5
    add_heading(doc, "3.5  Trade Journal (Transaction Records)", 2)
    add_heading(doc, "Log each buy/sell transaction", 3)
    for b in ["Stock ticker / name", "Direction: Buy or Sell", "Price, quantity, brokerage fees / commissions",
              "Date and time", "Notes (e.g. reason for entry)"]:
        add_bullet(doc, b, 1)
    add_heading(doc, "Open Positions Overview", 3)
    for b in [
        "Current holdings list with average cost, current price, unrealised P&L (amount + %)",
        "Total capital invested, total market value, total unrealised P&L",
    ]:
        add_bullet(doc, b, 1)
    add_bullet(doc, "Closed positions history (full trade log)")
    add_bullet(doc, "Monthly / yearly P&L summary chart")

    # 3.6 — NEW SECTION
    add_heading(doc, "3.6  Portfolio Performance Review  ★ NEW", 2)
    add_note(doc, "This module provides comprehensive analysis of the user's overall investment portfolio — both current and historical — to help users understand how their investments are performing and identify areas for improvement.")

    # 3.6.1
    add_heading(doc, "3.6.1  Overall Portfolio Summary", 3)
    for b in [
        "Total portfolio value (cost basis vs. current market value)",
        "Total realised P&L (from all closed trades)",
        "Total unrealised P&L (from all open positions)",
        "Overall portfolio return % (time-weighted)",
        "Cash balance / uninvested capital tracking",
    ]:
        add_bullet(doc, b, 1)

    # 3.6.2
    add_heading(doc, "3.6.2  Historical Performance Chart", 3)
    for b in [
        "Portfolio value over time — line chart (Daily / Monthly / Yearly view)",
        "Benchmark comparison: overlay HSI or HSTECH index for the same period",
        "Drawdown chart: visualise peak-to-trough declines over time",
    ]:
        add_bullet(doc, b, 1)

    # 3.6.3
    add_heading(doc, "3.6.3  Closed Trade Analysis", 3)
    for b in [
        "Complete list of all closed trades: entry price, exit price, hold duration, realised P&L",
        "Win rate: % of profitable trades vs. losing trades",
        "Average profit on winning trades vs. average loss on losing trades",
        "Profit factor: total gains ÷ total losses",
        "Best and worst individual trades (by absolute P&L and by %)",
        "Average holding period per trade",
    ]:
        add_bullet(doc, b, 1)

    # 3.6.4
    add_heading(doc, "3.6.4  Portfolio Composition & Allocation", 3)
    for b in [
        "Current asset allocation by sector (pie / donut chart)",
        "Concentration risk indicator: flag any single stock exceeding 20% of portfolio",
        "Top holdings table (ranked by current market value)",
    ]:
        add_bullet(doc, b, 1)

    # 3.6.5
    add_heading(doc, "3.6.5  Period Performance Reports", 3)
    for b in [
        "Selectable reporting period: weekly, monthly, quarterly, yearly, or custom date range",
        "For each period: opening value, closing value, net P&L, return %, number of trades",
        "Month-by-month performance calendar view (green / red tiles for positive / negative months)",
        "Exportable performance report (PDF or CSV)",
    ]:
        add_bullet(doc, b, 1)

    # 3.6.6
    add_heading(doc, "3.6.6  Dividend Income Tracking", 3)
    for b in [
        "Log dividends received per stock",
        "Total dividend income by month / year",
        "Projected annual dividend income based on current holdings and historical yield",
    ]:
        add_bullet(doc, b, 1)

    # 3.6.7
    add_heading(doc, "3.6.7  Risk Metrics", 3)
    for b in [
        "Portfolio Beta (sensitivity relative to HSI benchmark)",
        "Sharpe Ratio (risk-adjusted return)",
        "Volatility (annualised standard deviation of portfolio returns)",
        "Maximum Drawdown for the selected period",
    ]:
        add_bullet(doc, b, 1)

    # 3.7
    add_heading(doc, "3.7  Daily Market Digest", 2)
    for b in [
        "Push notification after each trading day closes",
        "Summary of overall market performance (indices)",
        "Performance of watchlist stocks",
        "Performance of current holdings",
        "Key financial news summary",
        "User-configurable push time (default: 4:30 PM HKT, after HKEX market close)",
    ]:
        add_bullet(doc, b)

    # 3.8
    add_heading(doc, "3.8  Analysis Tools", 2)
    for bold, text in [
        ("Stock Comparison Tool: ", "Compare up to 3 stocks side-by-side (price trend, key metrics)"),
        ("Position Sizing Calculator: ", "Calculate recommended position size based on account size, risk %, and stop-loss level"),
        ("Dividend Calculator: ", "Estimate projected annual dividend income from a given holding"),
    ]:
        add_bullet(doc, text, bold_prefix=bold)

    # 3.9
    add_heading(doc, "3.9  Settings", 2)
    for bold, text in [
        ("Language: ", "English (default) / Traditional Chinese"),
        ("Currency: ", "HKD (Hong Kong Dollar)"),
        ("Push Notifications: ", "On/Off toggle, configurable time"),
        ("Data Refresh Rate: ", "15 seconds / 1 minute / 5 minutes"),
        ("Export: ", "Trade journal and portfolio data as CSV"),
        ("Theme: ", "Dark mode (default) / Light mode"),
        ("Benchmark Index: ", "Select default benchmark (HSI / HSTECH / HSCEI)"),
    ]:
        add_bullet(doc, text, bold_prefix=bold)

    add_divider(doc)

    # ── 4. Non-Functional Requirements ──
    add_heading(doc, "4. Non-Functional Requirements", 1)
    nfr = [
        ("Category", "Requirement"),
        ("Performance", "Market data delay ≤ 15 seconds (free tier); charts load within 2 seconds"),
        ("Usability", "Clean and intuitive UI; common tasks completable within 3 taps"),
        ("Reliability", "Crash rate < 1%; offline cached data viewable when network is unavailable"),
        ("Security", "Trade records stored with local encryption; no account login required"),
        ("Compatibility", "Android 8.0 and above"),
        ("Localisation", "English default; Traditional Chinese switchable; HKT (UTC+8) timezone"),
        ("Data Integrity", "P&L calculations verified accurate; portfolio value reconciles with trade log"),
    ]
    t2 = doc.add_table(rows=len(nfr), cols=2)
    for i, (a, b) in enumerate(nfr):
        t2.cell(i, 0).text = a
        t2.cell(i, 1).text = b
    style_table(t2)
    add_divider(doc)

    # ── 5. Data Sources ──
    add_heading(doc, "5. Data Source Requirements", 1)
    ds = [
        ("Data Type", "Recommended Source"),
        ("Real-time / Delayed Quotes", "Yahoo Finance API, Alpha Vantage, Futu OpenAPI"),
        ("Historical Price Data", "Yahoo Finance (yfinance Python library)"),
        ("Financial News", "NewsAPI, AASTOCKS, HKEJ RSS feeds"),
        ("Fundamental Data", "Yahoo Finance, Macrotrends"),
        ("Sector Classification", "HKEX official website data"),
        ("Dividend Data", "Yahoo Finance, manual user input"),
        ("Benchmark Index Data", "Yahoo Finance (^HSI, ^HSCE, ^HSTECH)"),
    ]
    t3 = doc.add_table(rows=len(ds), cols=2)
    for i, (a, b) in enumerate(ds):
        t3.cell(i, 0).text = a
        t3.cell(i, 1).text = b
    style_table(t3)
    add_divider(doc)

    # ── 6. UI/UX ──
    add_heading(doc, "6. UI / UX Design Requirements", 1)
    for bold, text in [
        ("Style: ", "Modern financial app aesthetic (reference: Futu NiuNiu, Tiger Brokers)"),
        ("Default Theme: ", "Dark theme (deep navy / black background)"),
        ("Colour Convention: ", "Red = price up (gain);  Green = price down (loss) — Hong Kong / China market convention"),
        ("Charts: ", "Smooth, interactive candlestick charts with pinch-to-zoom and swipe"),
        ("Language: ", "All UI labels, menus, and notifications in English by default; fully translated to Traditional Chinese on language switch"),
    ]:
        add_bullet(doc, text, bold_prefix=bold)

    add_heading(doc, "Bottom Navigation Bar (5 tabs)", 3)
    for bold, text in [
        ("1. Market  ", "— Indices, gainers/losers, sector heatmap"),
        ("2. Watchlist  ", "— Personal watchlist + stock detail"),
        ("3. Portfolio  ", "— Open positions, trade journal, performance review"),
        ("4. Tools  ", "— Stock comparison, calculators"),
        ("5. Settings  ", "— Language, theme, notifications, export"),
    ]:
        add_bullet(doc, text, bold_prefix=bold)

    add_divider(doc)

    # ── 7. Navigation Structure ──
    add_heading(doc, "7. App Navigation Structure", 1)
    nav = [
        ("Tab", "Section", "Sub-pages / Features"),
        ("Market", "Indices Overview", "HSI, HSCEI, HSTECH live data"),
        ("", "Top Gainers / Losers", "Daily % change ranking"),
        ("", "Most Active", "By turnover value"),
        ("", "Sector Heatmap", "Industry-level performance"),
        ("Watchlist", "My Watchlist", "Grouped list with real-time quotes"),
        ("", "Stock Detail", "Chart, indicators, fundamentals, AI summary, news"),
        ("Portfolio", "Open Positions", "Holdings with unrealised P&L"),
        ("", "Trade Journal", "Log buy/sell transactions"),
        ("", "Performance Review", "Summary, charts, closed trade analysis, allocation, period reports, dividends, risk metrics"),
        ("", "Export Data", "CSV / PDF export"),
        ("Tools", "Stock Comparison", "Up to 3 stocks side-by-side"),
        ("", "Position Sizing Calc.", "Risk-based position calculator"),
        ("", "Dividend Calculator", "Projected income estimator"),
        ("Settings", "Language / Theme", "EN / 繁中, Dark / Light"),
        ("", "Notifications", "Push time and toggle"),
        ("", "Backup / Export", "Local data export"),
    ]
    t4 = doc.add_table(rows=len(nav), cols=3)
    for i, row in enumerate(nav):
        for j, val in enumerate(row):
            t4.cell(i, j).text = val
    style_table(t4)
    add_divider(doc)

    # ── 8. V2.0 ──
    add_heading(doc, "8. Future Version Features (V2.0 Roadmap)", 1)
    for b in [
        "Connect to broker API for automated trade execution (Futu / Tiger Brokers)",
        "AI stock screening and recommendation engine (technical + fundamental scoring)",
        "Portfolio stress testing and scenario analysis",
        "Tax report generation (for realised capital gains)",
        "Social / community discussion features",
        "iPad / iOS version",
    ]:
        add_bullet(doc, b)

    add_divider(doc)

    # ── 9. Change Log ──
    add_heading(doc, "9. Document Change Log", 1)
    changelog = [
        ("Version", "Date", "Changes"),
        ("v1.0", "2026-03-29", "Initial draft"),
        ("v1.1", "2026-03-29",
         "Changed default language to English (Traditional Chinese as switchable option); "
         "Added Portfolio Performance Review module (§3.6) with 7 sub-sections: Overall Summary, "
         "Historical Performance Chart, Closed Trade Analysis, Allocation Breakdown, Period Reports, "
         "Dividend Tracker, Risk Metrics; Added App Navigation Structure (§7); "
         "Updated Settings (§3.9) and UI/UX (§6) sections accordingly"),
    ]
    t5 = doc.add_table(rows=len(changelog), cols=3)
    for i, row in enumerate(changelog):
        for j, val in enumerate(row):
            t5.cell(i, j).text = val
    style_table(t5)

    path = os.path.join(BASE, "01_User_Requirements.docx")
    doc.save(path)
    print(f"Saved: {path}")

if __name__ == "__main__":
    build_user_requirements()
    print("Done!")
